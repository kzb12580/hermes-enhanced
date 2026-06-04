"""
Hermes Desktop Office 工具集 v2 — 真正的办公助手
增强: Excel 编辑/图表、PPT 图表/主题/模板、Word 高级排版
"""
import os
import io
import json
import logging
import shutil
import subprocess
from typing import Optional
from pathlib import Path

_log = logging.getLogger(__name__)

MAX_CONTENT_LEN = 10_000_000
MAX_ROWS = 500_000


def _safe_path(p: str) -> str:
    """路径净化 + 白名单保护"""
    resolved = str(Path(p).resolve())
    # 禁止写入系统目录
    blocked = ('/etc', '/usr', '/bin', '/sbin', '/boot', '/dev', '/proc', '/sys')
    if any(resolved.startswith(d) for d in blocked):
        raise ValueError(f"不允许写入系统路径: {resolved}")
    return resolved


def _check_file(p: str, name: str = "file") -> Optional[dict]:
    if p and not os.path.isfile(p):
        return {"error": f"{name} not found: {p}", "success": False}
    return None


# ═══════════════════════════════════════════════════════════════════════════
# Word 工具
# ═══════════════════════════════════════════════════════════════════════════

def create_word(path: str, title: str = "", content: str = "", template: str = "",
                font_size: int = 12, line_spacing: float = 1.5) -> dict:
    """创建 Word 文档（支持模板、字体、行距）"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if template:
            template = _safe_path(template)
            err = _check_file(template, "template")
            if err: return err
            doc = Document(template)
        else:
            doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(font_size)
        pf = style.paragraph_format
        pf.line_spacing = line_spacing

        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        path = _safe_path(path)
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        return {"path": path, "success": True}
    except Exception as e:
        _log.error("create_word failed: %s", e, exc_info=True)
        return {"error": str(e), "success": False}


def edit_word(path: str, operations: list[dict]) -> dict:
    """编辑 Word 文档"""
    try:
        path = _safe_path(path)
        err = _check_file(path, "document")
        if err: return err
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not isinstance(operations, list):
            return {"error": "operations must be a list", "success": False}
        doc = Document(path)

        for i, op in enumerate(operations):
            t = op.get("type")
            if t is None:
                return {"error": f"Op #{i} missing 'type'", "success": False}

            if t == "add_heading":
                level = max(0, min(op.get("level", 1), 9))
                h = doc.add_heading(op.get("text", ""), level=level)
                if op.get("align") == "center":
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "add_paragraph":
                p = doc.add_paragraph(op.get("text", ""))
                if op.get("bold"):
                    for run in p.runs:
                        run.bold = True
                if op.get("font_size"):
                    for run in p.runs:
                        run.font.size = Pt(op["font_size"])
                if op.get("align") == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif op.get("align") == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif t == "insert":
                idx = op.get("index", len(doc.paragraphs))
                text = op.get("text", "")
                if idx < len(doc.paragraphs):
                    doc.paragraphs[idx].insert_paragraph_before(text)
                else:
                    doc.add_paragraph(text)
            elif t == "delete_paragraph":
                idx = op.get("index", -1)
                if 0 <= idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p._element.getparent().remove(p._element)
                else:
                    return {"error": f"Invalid paragraph index: {idx}", "success": False}
            elif t == "add_table":
                rows = op.get("rows", [])
                if not rows or not rows[0]:
                    return {"error": f"Op #{i}: empty table", "success": False}
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols, style=op.get("style", "Table Grid"))
                for r, row_data in enumerate(rows):
                    for c, cell_val in enumerate(row_data):
                        table.rows[r].cells[c].text = str(cell_val)
            elif t == "add_image":
                img_path = op.get("image_path", "")
                if img_path:
                    img_path = str(_safe_path(img_path))
                    if isinstance(img_path, str) and img_path.startswith(("Error", "⚠️")):
                        return {"error": f"Op #{i}: {img_path}", "success": False}
                    err = _check_file(img_path, "image")
                    if err: return err
                    doc.add_picture(img_path, width=Inches(op.get("width", 6)))
            elif t == "add_page_break":
                doc.add_page_break()
            elif t == "replace":
                old, new = op.get("old", ""), op.get("new", "")
                for p in doc.paragraphs:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
            elif t == "set_header":
                section = doc.sections[0]
                header = section.header
                header.paragraphs[0].text = op.get("text", "")
            elif t == "set_footer":
                section = doc.sections[0]
                footer = section.footer
                footer.paragraphs[0].text = op.get("text", "")
            elif t == "add_toc":
                # 添加目录域 (TOC)
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                # begin
                fldChar_begin = OxmlElement('w:fldChar')
                fldChar_begin.set(qn('w:fldCharType'), 'begin')
                run._r.append(fldChar_begin)
                # instrText
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
                run._r.append(instrText)
                # separate
                fldChar_sep = OxmlElement('w:fldChar')
                fldChar_sep.set(qn('w:fldCharType'), 'separate')
                run._r.append(fldChar_sep)
                # placeholder
                run2 = paragraph.add_run("(请在Word中右键更新目录)")
                # end
                fldChar_end = OxmlElement('w:fldChar')
                fldChar_end.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar_end)
            else:
                return {"error": f"Unknown op type: {t}", "success": False}

        doc.save(path)
        return {"path": path, "operations": len(operations), "success": True}
    except Exception as e:
        _log.error("edit_word failed: %s", e, exc_info=True)
        return {"error": str(e), "success": False}


def read_word(path: str) -> dict:
    """读取 Word 文档内容"""
    try:
        path = _safe_path(path)
        err = _check_file(path, "document")
        if err: return err
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return {"paragraphs": paragraphs, "tables": tables, "success": True}
    except Exception as e:
        return {"error": str(e), "success": False}


# ═══════════════════════════════════════════════════════════════════════════
# PPT 工具 — 基于 PptxGenJS (Node.js)
# ═══════════════════════════════════════════════════════════════════════════


def create_ppt(path: str, slides: Optional[list[dict]] = None, layout: str = "16x9",
               title: str = "", author: str = "", slides_file: str = "") -> dict:
    """
    创建 PPT 演示文稿（基于 PptxGenJS，支持动画/过渡/阴影/透明度）

    参数:
      - path: 输出文件路径
      - slides: 幻灯片数组（简单PPT直接传，复杂PPT用slides_file）
      - slides_file: 幻灯片JSON文件路径（避免大JSON截断，推荐>5页时使用）
      - layout: 布局 (16x9 / 16x10 / 4x3 / wide)
      - title: 演示文稿标题
      - author: 作者

    ⚠️ 重要：>5页PPT请用 slides_file 方式避免截断！
    步骤：1) write_file("slides.json", JSON数组) → 2) create_ppt(path, slides_file="slides.json")

    每个 slide 的结构:
      {
        "background": {"color": "1E2761"},
        "transition": {"type": "fade", "duration": 1},
        "elements": [
          {"type": "text", "text": "标题", "x": 1, "y": 2, "w": 8, "h": 2,
           "fontSize": 44, "bold": true, "color": "FFFFFF", "align": "center"},
          {"type": "shape", "shape": "rect", "x": 0, "y": 6.5, "w": 13.333, "h": 1,
           "fill": {"color": "065A82"}},
          {"type": "image", "path": "https://...", "x": 1, "y": 3, "w": 5, "h": 3},
          {"type": "chart", "chartType": "bar",
           "data": [{"name": "销量", "labels": ["Q1","Q2"], "values": [450,550]}],
           "x": 0.5, "y": 1, "w": 9, "h": 4},
          {"type": "table", "rows": [["H1","H2"],["c1","c2"]], "x": 1, "y": 1, "w": 8, "h": 2}
        ]
      }

    text 元素支持富文本数组: "text": [{"text": "粗体", "bold": true}, "普通"]

    支持的 chartType: bar, line, pie, doughnut, scatter, radar
    支持的 shape: rect, oval, line, rounded_rect
    支持的 transition type: fade, push, cover, uncover, wipe, split, blinds, checkerboard, random
    """
    try:
        # slides_file 优先：从JSON文件读取slides（避免大JSON截断）
        if slides_file:
            sf = _safe_path(slides_file)
            if not os.path.isfile(sf):
                return {"error": f"slides_file 不存在: {sf}", "success": False}
            try:
                with open(sf, "r", encoding="utf-8-sig") as f:  # utf-8-sig 自动strip BOM
                    slides = json.load(f)
            except (json.JSONDecodeError, IOError) as e:
                return {"error": f"slides_file 读取失败: {e}", "success": False}

        if not slides or not isinstance(slides, list):
            return {"error": "slides 参数缺失或为空。请提供 slides 数组或 slides_file 路径。", "success": False}

        path = _safe_path(path)
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)

        config = {
            "path": os.path.abspath(path),
            "layout": layout,
            "slides": slides,
        }
        if title:
            config["title"] = title
        if author:
            config["author"] = author

        # Ensure pptxgenjs is installed (auto-install on first use)
        tools_dir = os.path.dirname(os.path.abspath(__file__))
        node_modules = os.path.join(tools_dir, "node_modules", "pptxgenjs")
        if not os.path.isdir(node_modules):
            _log.info("pptxgenjs not found, installing...")
            install_result = subprocess.run(
                ["npm", "install", "--production"],
                cwd=tools_dir, capture_output=True, text=True, timeout=120, encoding="utf-8",
            )
            if install_result.returncode != 0:
                return {"error": f"Failed to install pptxgenjs: {install_result.stderr.strip()}", "success": False}

        # Write config to temp file, call Node.js worker
        import tempfile
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False)
            config_path = f.name

        try:
            worker_script = os.path.join(os.path.dirname(__file__), "pptxgenjs_worker.js")
            result = subprocess.run(
                ["node", worker_script],
                input=json.dumps(config, ensure_ascii=False),
                capture_output=True, text=True, timeout=120, encoding="utf-8",
            )
            if result.returncode != 0:
                error_msg = result.stderr.strip() or result.stdout.strip() or "Unknown error"
                return {"error": f"PptxGenJS error: {error_msg}", "success": False}

            output = json.loads(result.stdout)
            return output
        finally:
            os.unlink(config_path)

    except json.JSONDecodeError as e:
        _log.error("create_ppt JSON parse error: %s", e)
        return {"error": f"Failed to parse PptxGenJS output: {e}", "success": False}
    except subprocess.TimeoutExpired:
        return {"error": "PPT generation timed out (120s)", "success": False}
    except Exception as e:
        _log.error("create_ppt failed: %s", e, exc_info=True)
        return {"error": str(e), "success": False}


# ═══════════════════════════════════════════════════════════════════════════
# Excel 工具 — 增强版：编辑现有文件、图表、公式、格式
# ═══════════════════════════════════════════════════════════════════════════

def create_excel(path: str, sheets: list[dict]) -> dict:
    """创建 Excel 表格"""
    try:
        total_rows = sum(len(s.get("data", [])) for s in sheets)
        if total_rows > MAX_ROWS:
            return {"error": f"Too many rows ({total_rows})", "success": False}

        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()
        if not isinstance(sheets, list) or not sheets:
            return {"error": "sheets must be a non-empty list", "success": False}

        for i, sheet_data in enumerate(sheets):
            ws = wb.active if i == 0 else wb.create_sheet()
            ws.title = sheet_data.get("name", f"Sheet{i+1}")

            rows = sheet_data.get("data", [])
            headers = sheet_data.get("headers", [])

            if headers:
                for j, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=j, value=header)
                    cell.font = Font(color="FFFFFF", bold=True, size=11)
                    cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                start_row = 2
            else:
                start_row = 1

            for r, row_data in enumerate(rows, start_row):
                for c, value in enumerate(row_data, 1):
                    cell = ws.cell(row=r, column=c, value=value)
                    cell.alignment = Alignment(vertical="center")

            # 自动列宽
            for col in ws.columns:
                max_len = 0
                col_letter = col[0].column_letter
                for cell in col:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

        path = _safe_path(path)
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        wb.save(path)
        return {"path": path, "sheets": len(sheets), "success": True}
    except Exception as e:
        _log.error("create_excel failed: %s", e, exc_info=True)
        return {"error": str(e), "success": False}


def read_excel(path: str, sheet_name: str = "", max_rows: int = 1000) -> dict:
    """读取 Excel 文件"""
    try:
        path = _safe_path(path)
        err = _check_file(path, "excel")
        if err: return err
        # 文件大小检查
        file_size = os.path.getsize(path)
        if file_size > 50_000_000:  # 50MB
            return {"error": f"Excel文件过大 ({file_size // 1_000_000}MB > 50MB)", "success": False}
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        sheet_names_cache = list(wb.sheetnames)  # close前缓存
        result = {}
        sheets_to_read = [sheet_name] if sheet_name and sheet_name in sheet_names_cache else sheet_names_cache
        for name in sheets_to_read:
            ws = wb[name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                rows.append(list(row))
            result[name] = rows
        wb.close()
        return {"sheets": result, "sheet_names": sheet_names_cache, "success": True}
    except Exception as e:
        return {"error": str(e), "success": False}


def _safe_ref(ws, ref_dict):
    """Safe Reference builder for openpyxl charts."""
    from openpyxl.chart import Reference
    if not isinstance(ref_dict, dict):
        return Reference(ws, min_col=1, min_row=1, max_col=1, max_row=1)
    return Reference(ws,
        min_row=max(1, ref_dict.get("min_row", 1)),
        max_row=max(1, ref_dict.get("max_row", 1)),
        min_col=max(1, ref_dict.get("min_col", 1)),
        max_col=max(1, ref_dict.get("max_col", 1)))


def edit_excel(path: str, operations: list[dict]) -> dict:
    """
    编辑 Excel 文件

    operations 支持:
      - set_cell: 设置单元格 {row, col, value, sheet}
      - set_range: 批量设置 {start_row, start_col, data, sheet}
      - add_sheet: 添加工作表 {name}
      - delete_sheet: 删除工作表 {name}
      - add_chart: 添加图表 {sheet, chart_type, data_range, title, position}
      - add_formula: 添加公式 {row, col, formula, sheet}
      - format_cells: 格式化 {range, bold, color, bg_color, align, sheet}
      - auto_filter: 添加筛选 {sheet, range}
      - merge_cells: 合并单元格 {range, sheet}
    """
    try:
        path = _safe_path(path)
        err = _check_file(path, "excel")
        if err: return err
        from openpyxl import load_workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        from openpyxl.utils import get_column_letter

        wb = load_workbook(path)
        if not isinstance(operations, list):
            return {"error": "operations must be a list", "success": False}

        for i, op in enumerate(operations):
            t = op.get("type")
            if t is None:
                return {"error": f"Op #{i} missing 'type'", "success": False}

            sheet_name = op.get("sheet", wb.sheetnames[0])
            if sheet_name not in wb.sheetnames:
                return {"error": f"Sheet '{sheet_name}' not found", "success": False}
            ws = wb[sheet_name]

            if t == "set_cell":
                ws.cell(row=op["row"], column=op["col"], value=op.get("value"))

            elif t == "set_range":
                data = op.get("data", [])
                sr, sc = op.get("start_row", 1), op.get("start_col", 1)
                for r, row_data in enumerate(data):
                    for c, val in enumerate(row_data):
                        ws.cell(row=sr + r, column=sc + c, value=val)

            elif t == "add_sheet":
                wb.create_sheet(op.get("name", f"Sheet{len(wb.sheetnames)+1}"))

            elif t == "delete_sheet":
                if op.get("name") in wb.sheetnames and len(wb.sheetnames) > 1:
                    del wb[op["name"]]

            elif t == "add_chart":
                chart_types = {"bar": BarChart, "line": LineChart, "pie": PieChart}
                chart_cls = chart_types.get(op.get("chart_type", "bar"), BarChart)
                chart = chart_cls()
                chart.title = op.get("title", "")
                chart.style = 10
                if chart_cls is not PieChart:
                    chart.y_axis.title = op.get("y_axis", "")
                    chart.x_axis.title = op.get("x_axis", "")
                data_ref = _safe_ref(ws, op.get("data_ref", {}))
                cats_ref = _safe_ref(ws, op.get("cats_ref", {}))
                chart.add_data(data_ref, titles_from_data=True)
                chart.set_categories(cats_ref)
                chart.width = op.get("width", 20)
                chart.height = op.get("height", 12)
                pos = op.get("position", "E2")
                ws.add_chart(chart, pos)

            elif t == "add_formula":
                ws.cell(row=op["row"], column=op["col"], value=op.get("formula", ""))

            elif t == "format_cells":
                cell_range = op.get("range", "A1:A1")
                bold = op.get("bold", False)
                font_color = op.get("color")
                bg_color = op.get("bg_color")
                font_size = op.get("font_size", 11)
                for row in ws[cell_range]:
                    for cell in row:
                        if bold or font_color or font_size:
                            cell.font = Font(bold=bold, color=font_color, size=font_size)
                        if bg_color:
                            cell.fill = PatternFill(start_color=bg_color, end_color=bg_color, fill_type="solid")
                        if op.get("align"):
                            cell.alignment = Alignment(horizontal=op["align"])

            elif t == "auto_filter":
                ws.auto_filter.ref = op.get("range", ws.dimensions)

            elif t == "merge_cells":
                ws.merge_cells(op.get("range", "A1:B1"))

            else:
                return {"error": f"Unknown op type: {t}", "success": False}

        wb.save(path)
        return {"path": path, "operations": len(operations), "success": True}
    except Exception as e:
        _log.error("edit_excel failed: %s", e, exc_info=True)
        return {"error": str(e), "success": False}


# ═══════════════════════════════════════════════════════════════════════════
# 工具注册 — 通过 office_tool_wrappers.py 的 OFFICE_TOOL_DEFINITIONS 注册
# ═══════════════════════════════════════════════════════════════════════════
